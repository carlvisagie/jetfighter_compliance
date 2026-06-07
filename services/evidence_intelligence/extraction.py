"""Text extraction from uploaded evidence (rule-based v1)."""
from __future__ import annotations

import csv
import io
import json
import re
from pathlib import Path
from typing import Tuple

from .ocr import (
    DEFAULT_OCR_MAX_PAGES,
    looks_like_scanned_pdf,
    ocr_enabled,
    ocr_image_bytes,
    ocr_pdf_bytes,
)
from .schemas import ExtractionResult

MAX_SYNC_BYTES = 2 * 1024 * 1024
MAX_PREVIEW = 8000
MAX_SNIPPET_STORE = 240

SECRET_PATTERNS = [
    re.compile(r"(?i)(api[_-]?key|secret[_-]?key?|password|passwd|token|bearer|credential|auth[_-]?key)\s*[:=]\s*\S{8,}"),
    re.compile(r"AKIA[0-9A-Z]{16}"),                                          # AWS access key
    re.compile(r"ASIA[0-9A-Z]{16}"),                                          # AWS temp key
    re.compile(r"-----BEGIN (?:RSA |EC |DSA |OPENSSH |PGP )?PRIVATE KEY(?: BLOCK)?-----"),
    re.compile(r"eyJ[A-Za-z0-9_-]{10,}\.[A-Za-z0-9_-]{10,}"),               # JWT
    re.compile(r"ghp_[A-Za-z0-9]{36}"),                                       # GitHub PAT
    re.compile(r"sk-[a-zA-Z0-9]{32,}"),                                       # OpenAI key pattern
    re.compile(r"(?i)private[_\s-]?key\s*[:=]\s*\S{16,}"),
]

DANGEROUS_EXTENSIONS = {".exe", ".bat", ".cmd", ".ps1", ".sh", ".dll", ".msi", ".scr", ".vbs", ".js"}


def redact_secrets(text: str) -> str:
    if not text:
        return text
    out = text
    for rx in SECRET_PATTERNS:
        out = rx.sub("[REDACTED]", out)
    return out


def safe_snippet(text: str, limit: int = MAX_SNIPPET_STORE) -> str:
    t = redact_secrets((text or "").strip())
    if len(t) > limit:
        return t[: limit - 3] + "..."
    return t


def _extract_txt(data: bytes) -> Tuple[str, str]:
    for enc in ("utf-8", "latin-1"):
        try:
            return data.decode(enc), "text_plain"
        except UnicodeDecodeError:
            continue
    return "", "decode_failed"


def _extract_pdf(data: bytes) -> Tuple[str, str]:
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        parts = []
        for page in reader.pages[:30]:
            parts.append(page.extract_text() or "")
        return "\n".join(parts), "pdf_text"
    except ImportError:
        return "", "pdf_unavailable"
    except Exception:
        return "", "pdf_failed"


def _extract_csv(data: bytes) -> Tuple[str, str]:
    text, _ = _extract_txt(data)
    if not text:
        return "", "csv_empty"
    try:
        rows = list(csv.reader(io.StringIO(text)))
        return "\n".join(", ".join(r) for r in rows[:200]), "csv_parse"
    except Exception:
        return text[:MAX_PREVIEW], "csv_raw"


def _extract_docx(data: bytes) -> Tuple[str, str]:
    try:
        from docx import Document
    except ImportError:
        return "", "docx_unavailable"
    try:
        doc = Document(io.BytesIO(data))
        parts: list[str] = []
        for para in doc.paragraphs:
            t = (para.text or "").strip()
            if t:
                parts.append(t)
            if sum(len(p) for p in parts) > MAX_PREVIEW * 4:
                break
        for table in doc.tables:
            for row in table.rows:
                cells = [(c.text or "").strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
                if sum(len(p) for p in parts) > MAX_PREVIEW * 4:
                    break
        return "\n".join(parts), "docx_text"
    except Exception:
        return "", "docx_failed"


def _extract_xlsx(data: bytes) -> Tuple[str, str]:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "", "xlsx_unavailable"
    try:
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts: list[str] = []
        for sheet in wb.worksheets:
            parts.append(f"# Sheet: {sheet.title}")
            count = 0
            for row in sheet.iter_rows(values_only=True):
                cells = [
                    (str(v).strip() if v is not None else "") for v in row
                ]
                if any(cells):
                    parts.append(" | ".join(cells))
                    count += 1
                if count >= 200:
                    break
            if sum(len(p) for p in parts) > MAX_PREVIEW * 4:
                break
        wb.close()
        return "\n".join(parts), "xlsx_text"
    except Exception:
        return "", "xlsx_failed"


def _extract_image_metadata(data: bytes, ext: str) -> Tuple[str, str]:
    """Surface image metadata as a textual summary.

    No OCR (deferred — requires tesseract system binary). What we extract here
    is enough to (a) prove the upload is a real image, (b) route it for human
    review, (c) feed classification signals (filename + dimensions + format).
    """
    try:
        from PIL import Image, ExifTags
    except ImportError:
        return "", "image_metadata_unavailable"
    try:
        img = Image.open(io.BytesIO(data))
        img.load()
        width, height = img.size
        mode = img.mode
        fmt = img.format or ext.lstrip(".").upper()
        exif_lines: list[str] = []
        try:
            exif = img.getexif() or {}
            for tag_id, value in exif.items():
                tag_name = ExifTags.TAGS.get(tag_id, str(tag_id))
                if isinstance(value, bytes):
                    continue
                v = str(value).strip()
                if v and len(v) < 200:
                    exif_lines.append(f"{tag_name}: {v}")
                if len(exif_lines) >= 20:
                    break
        except Exception:
            pass
        summary = [
            f"image_format: {fmt}",
            f"dimensions: {width}x{height}",
            f"mode: {mode}",
        ]
        if exif_lines:
            summary.append("exif:")
            summary.extend(exif_lines)
        summary.append("note: OCR not enabled — text inside image awaits manual review.")
        return "\n".join(summary), "image_metadata"
    except Exception:
        return "", "image_metadata_failed"


def extract_from_file(path: Path, *, size_bytes: int = 0) -> ExtractionResult:
    name = path.name
    ext = path.suffix.lower()
    result = ExtractionResult(source_file=name, mime_hint=ext)

    if ext in DANGEROUS_EXTENSIONS:
        result.ok = False
        result.pending_analysis = True
        result.errors.append("unsupported_file_type")
        result.extraction_method = "rejected_type"
        return result

    size = size_bytes or path.stat().st_size
    if size > MAX_SYNC_BYTES:
        result.pending_analysis = True
        result.warnings.append("file_too_large_for_sync_analysis")
        result.extraction_method = "pending"
        return result

    try:
        data = path.read_bytes()
    except Exception as e:
        result.ok = False
        result.errors.append(str(e)[:120])
        return result

    text = ""
    method = "none"
    if ext in (".txt", ".md", ".log", ".json", ".xml", ".html", ".htm"):
        text, method = _extract_txt(data)
        if ext == ".json":
            method = "json_text"
    elif ext == ".pdf":
        text, method = _extract_pdf(data)
        if method == "pdf_unavailable":
            result.warnings.append("pdf_extraction_unavailable")
            result.pending_analysis = True
        # OCR fallback for scanned PDFs: pypdf returns only whitespace
        # for image-only PDFs. If the extracted text is below a
        # conservative threshold *and* OCR is enabled, try OCR.
        if ocr_enabled() and looks_like_scanned_pdf(text):
            ocr_text, ocr_status = ocr_pdf_bytes(
                data, max_pages=DEFAULT_OCR_MAX_PAGES
            )
            result.ocr_status = ocr_status
            if ocr_status == "ok" or ocr_status == "ocr_ok":
                result.ocr_applied = True
                result.ocr_text_length = len(ocr_text)
                # Preserve any text pypdf did find by prepending it.
                if text.strip():
                    text = text + "\n\n--- ocr text ---\n" + ocr_text
                else:
                    text = ocr_text
                method = "pdf_ocr"
                # Clear pending if we now have real text.
                if text.strip():
                    result.pending_analysis = False
            elif ocr_status in (
                "ocr_module_unavailable",
                "ocr_binary_unavailable",
            ):
                result.warnings.append(
                    f"scanned_pdf_ocr_skipped:{ocr_status}"
                )
                result.pending_analysis = True
            elif ocr_status == "ocr_empty":
                result.warnings.append("scanned_pdf_ocr_empty")
                result.pending_analysis = True
            elif ocr_status == "ocr_failed":
                result.warnings.append("scanned_pdf_ocr_failed")
                result.pending_analysis = True
            elif ocr_status == "ocr_timeout":
                result.warnings.append("scanned_pdf_ocr_timeout")
                result.pending_analysis = True
            elif ocr_status == "ocr_corrupted_pdf":
                result.warnings.append("scanned_pdf_ocr_corrupted")
                result.pending_analysis = True
    elif ext in (".csv",):
        text, method = _extract_csv(data)
    elif ext == ".docx":
        text, method = _extract_docx(data)
        if method == "docx_unavailable":
            result.warnings.append("docx_extraction_unavailable")
            result.pending_analysis = True
        elif method == "docx_failed":
            result.warnings.append("docx_extraction_failed")
            result.pending_analysis = True
    elif ext == ".xlsx":
        text, method = _extract_xlsx(data)
        if method == "xlsx_unavailable":
            result.warnings.append("xlsx_extraction_unavailable")
            result.pending_analysis = True
        elif method == "xlsx_failed":
            result.warnings.append("xlsx_extraction_failed")
            result.pending_analysis = True
    elif ext in (".doc", ".xls"):
        # Legacy binary Office formats — not in scope for v1 extraction.
        # Customer should re-export as .docx/.xlsx; flag for human review.
        result.pending_analysis = True
        result.warnings.append("legacy_office_binary_format")
        method = "office_legacy_pending"
    elif ext in (".zip", ".7z", ".rar", ".tar", ".gz", ".tgz"):
        # Archives must never fall through to text extraction — that
        # path returns latin-1 garbage which then pollutes classify /
        # entity steps. Production intake FB-1dfab13c120b uploaded a
        # 4.7 MB packet.zip that was silently classified as plain text
        # before this branch existed. Flag for manual review with an
        # actionable operator hint so the customer is asked to re-
        # upload the contents individually.
        result.pending_analysis = True
        result.warnings.append("archive_pending_manual_extraction")
        method = "archive_pending"
    elif ext in (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff", ".tif"):
        text, method = _extract_image_metadata(data, ext)
        # Image METADATA is real text; mark pending only if metadata extraction
        # itself failed. The image still goes through classification + entity
        # extraction on filename + metadata.
        if method in ("image_metadata_unavailable", "image_metadata_failed"):
            result.pending_analysis = True
        # OCR on the image itself if enabled — this is what surfaces
        # MFA screenshots, scanned licences, smartphone-photo policies.
        if ocr_enabled():
            ocr_text, ocr_status = ocr_image_bytes(data)
            result.ocr_status = ocr_status
            if ocr_status == "ocr_ok":
                result.ocr_applied = True
                result.ocr_text_length = len(ocr_text)
                text = (text + "\n\n--- ocr text ---\n" + ocr_text) if text else ocr_text
                method = "image_ocr"
                result.pending_analysis = False
            elif ocr_status in (
                "ocr_module_unavailable",
                "ocr_binary_unavailable",
            ):
                result.warnings.append(
                    f"image_ocr_skipped:{ocr_status}"
                )
            elif ocr_status == "ocr_empty":
                result.warnings.append("image_ocr_empty")
            elif ocr_status == "ocr_failed":
                result.warnings.append("image_ocr_failed")
        else:
            result.warnings.append("image_text_awaits_ocr_for_full_extraction")
    else:
        text, method = _extract_txt(data)
        if not text.strip():
            result.pending_analysis = True
            result.warnings.append("unknown_format")

    text = redact_secrets(text)
    result.text_length = len(text)
    result.text_preview = safe_snippet(text, MAX_PREVIEW)
    result.extraction_method = method
    if not text.strip() and not result.pending_analysis:
        result.warnings.append("no_extractable_text")
    return result
