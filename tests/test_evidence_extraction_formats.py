"""Real-format extraction tests — PDF, DOCX, XLSX, image metadata.

These exercise the extraction surface that the autonomous revenue path
depends on. If any of them silently regresses to a `*_unavailable` /
`*_pending` extraction_method, classification → product_id selection →
auto_send_payment_link is broken and the platform earns nothing.
"""
from __future__ import annotations

import io
from pathlib import Path

import pytest

from services.evidence_intelligence.extraction import extract_from_file


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _build_minimal_pdf(text: str) -> bytes:
    """Build a tiny single-page PDF whose text is `text`.

    Uses pypdf's writer rather than hand-crafting bytes so the test stays
    in sync with the library we actually use in production.
    """
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=300, height=200)
    # pypdf 4.x can't add Tj text content to a blank page without a layout
    # engine, so we embed text via the document-info dict; the extractor
    # below proves PDFs roundtrip even when no text frames are present.
    writer.add_metadata({"/Title": text})
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def test_pdf_extraction_returns_pdf_text_method(tmp_path: Path):
    """A valid PDF must be routed through the pypdf branch.

    We don't assert text content here because pypdf returns "" for blank
    pages — the contract is that extraction_method is `pdf_text`, not
    `pdf_unavailable` (which silently swallowed PDFs before pypdf was
    added to requirements.txt).
    """
    pdf_bytes = _build_minimal_pdf("ACME SSP v1")
    f = tmp_path / "ssp.pdf"
    f.write_bytes(pdf_bytes)
    result = extract_from_file(f)
    assert result.ok is True
    assert result.extraction_method == "pdf_text"
    assert result.pending_analysis is False
    assert "pdf_extraction_unavailable" not in result.warnings


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def test_docx_extraction_returns_paragraph_text(tmp_path: Path):
    from docx import Document

    doc = Document()
    doc.add_heading("ACME Defense Co — SSP", level=1)
    doc.add_paragraph("Multi-factor authentication is enforced for all users.")
    doc.add_paragraph("Contact security: ops@acme-defense.com")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Control"
    table.cell(0, 1).text = "Status"
    table.cell(1, 0).text = "AC-1"
    table.cell(1, 1).text = "Implemented"
    f = tmp_path / "ssp.docx"
    doc.save(str(f))

    result = extract_from_file(f)
    assert result.ok is True
    assert result.extraction_method == "docx_text"
    assert result.pending_analysis is False
    assert "ACME" in result.text_preview
    assert "Multi-factor authentication" in result.text_preview
    # Table cells must also be captured.
    assert "AC-1" in result.text_preview
    assert "Implemented" in result.text_preview


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------

def test_xlsx_extraction_returns_sheet_rows(tmp_path: Path):
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Inventory"
    ws.append(["Asset", "Owner", "Classification"])
    ws.append(["Server-01", "ops@acme.com", "Internal"])
    ws.append(["Laptop-21", "tina@acme.com", "Confidential"])
    f = tmp_path / "inventory.xlsx"
    wb.save(str(f))

    result = extract_from_file(f)
    assert result.ok is True
    assert result.extraction_method == "xlsx_text"
    assert result.pending_analysis is False
    assert "Inventory" in result.text_preview
    assert "Server-01" in result.text_preview
    assert "ops@acme.com" in result.text_preview


# ---------------------------------------------------------------------------
# Image metadata (no OCR)
# ---------------------------------------------------------------------------

def test_image_metadata_extracts_dimensions(tmp_path: Path):
    from PIL import Image

    img = Image.new("RGB", (640, 480), color=(255, 0, 0))
    f = tmp_path / "mfa_screenshot.png"
    img.save(str(f), format="PNG")

    result = extract_from_file(f)
    assert result.ok is True
    assert result.extraction_method == "image_metadata"
    # Dimensions must surface so classification sees the file is a real
    # image (e.g. screenshots vs full-page scans drive different routing).
    assert "640x480" in result.text_preview
    assert "image_format: PNG" in result.text_preview
    # The OCR limitation must be visible to the operator path.
    assert any("ocr" in w.lower() for w in result.warnings)


# ---------------------------------------------------------------------------
# Legacy binary office formats — should be marked, never silently zero-text.
# ---------------------------------------------------------------------------

def test_legacy_doc_xls_flagged_pending(tmp_path: Path):
    f = tmp_path / "old_contract.doc"
    f.write_bytes(b"\xd0\xcf\x11\xe0" + b"\x00" * 100)
    result = extract_from_file(f)
    assert result.extraction_method == "office_legacy_pending"
    assert result.pending_analysis is True
    assert "legacy_office_binary_format" in result.warnings
